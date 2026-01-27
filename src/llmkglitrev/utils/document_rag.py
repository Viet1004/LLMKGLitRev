"""
Document RAG system for industry partner review.

Implements a simple FAISS-based RAG to load project documents (PDF, DOCX)
and enable LLMs to ask alignment questions based on document context.
"""

from pathlib import Path
from typing import List, Dict, Optional
import tempfile
import os

# Document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Embeddings and vector store
try:
    from langchain_openai import OpenAIEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False


class ProjectDocumentRAG:
    """
    Simple RAG system for project documents.

    Allows industry partners to ask questions about:
    - Research-project alignment
    - Implementation feasibility
    - Resource requirements
    - Timeline expectations
    """

    def __init__(self, session_id: str):
        """Initialize RAG system for a session."""
        self.session_id = session_id
        self.session_dir = Path(f"sessions/{session_id}")
        self.docs_dir = self.session_dir / "project_documents"
        self.docs_dir.mkdir(parents=True, exist_ok=True)

        self.vector_store: Optional[FAISS] = None
        self.documents: List[Document] = []

        if not LANGCHAIN_AVAILABLE:
            raise ImportError("LangChain not available. Install: pip install langchain langchain-openai langchain-community")

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install: pip install PyPDF2")

        text = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text())

        return "\n\n".join(text)

    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install: pip install python-docx")

        doc = DocxDocument(file_path)
        text = []

        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        return "\n\n".join(text)

    def add_document(self, file_path: Path, metadata: Optional[Dict] = None) -> None:
        """
        Add a document to the RAG system.

        Args:
            file_path: Path to PDF or DOCX file
            metadata: Optional metadata about the document
        """
        # Extract text based on file type
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif suffix == '.docx':
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Use .pdf or .docx")

        # Create document with metadata
        doc_metadata = metadata or {}
        doc_metadata.update({
            "source": file_path.name,
            "file_type": suffix
        })

        doc = Document(
            page_content=text,
            metadata=doc_metadata
        )

        self.documents.append(doc)

    def build_index(self, chunk_size: int = 1000, chunk_overlap: int = 200) -> None:
        """
        Build FAISS index from documents.

        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        if not self.documents:
            raise ValueError("No documents added. Call add_document() first.")

        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        chunks = text_splitter.split_documents(self.documents)

        # Create embeddings and vector store
        embeddings = OpenAIEmbeddings()
        self.vector_store = FAISS.from_documents(chunks, embeddings)

        # Save to disk
        index_path = self.docs_dir / "faiss_index"
        self.vector_store.save_local(str(index_path))

    def load_index(self) -> bool:
        """
        Load existing FAISS index from disk.

        Returns:
            True if index was loaded successfully, False otherwise
        """
        index_path = self.docs_dir / "faiss_index"

        if not index_path.exists():
            return False

        try:
            embeddings = OpenAIEmbeddings()
            self.vector_store = FAISS.load_local(
                str(index_path),
                embeddings,
                allow_dangerous_deserialization=True
            )
            return True
        except Exception as e:
            print(f"Error loading index: {e}")
            return False

    def retrieve_context(self, query: str, k: int = 3) -> List[Dict]:
        """
        Retrieve relevant context for a query.

        Args:
            query: Question or topic to retrieve context for
            k: Number of chunks to retrieve

        Returns:
            List of dicts with 'content' and 'metadata'
        """
        if not self.vector_store:
            raise ValueError("No vector store available. Build or load index first.")

        results = self.vector_store.similarity_search(query, k=k)

        return [
            {
                "content": doc.page_content,
                "metadata": doc.metadata
            }
            for doc in results
        ]

    def format_context_for_prompt(self, query: str, k: int = 3) -> str:
        """
        Format retrieved context for LLM prompt.

        Args:
            query: Question to retrieve context for
            k: Number of chunks to retrieve

        Returns:
            Formatted context string
        """
        contexts = self.retrieve_context(query, k=k)

        formatted = ["<Project Documentation Context>"]

        for i, ctx in enumerate(contexts, 1):
            source = ctx['metadata'].get('source', 'Unknown')
            formatted.append(f"\n[Document {i}: {source}]")
            formatted.append(ctx['content'])

        formatted.append("\n</Project Documentation Context>")

        return "\n".join(formatted)

    def get_document_summary(self) -> Dict:
        """Get summary of indexed documents."""
        return {
            "num_documents": len(self.documents),
            "documents": [
                {
                    "name": doc.metadata.get("source", "Unknown"),
                    "type": doc.metadata.get("file_type", "Unknown"),
                    "length": len(doc.page_content)
                }
                for doc in self.documents
            ],
            "indexed": self.vector_store is not None
        }
